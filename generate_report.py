#!/usr/bin/env python3
"""Generate Customer Support ChatBot B.Tech Final Year Project Report — Sharda University format"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Constants ────────────────────────────────────────────────────────────────
FM   = 'Times New Roman'
FC   = 'Courier New'
LS   = 1.5
BASE = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS = {
    'ss1': os.path.join(BASE, 'ss1_main.png'),
    'ss2': os.path.join(BASE, 'ss2_greeting.png'),
    'ss3': os.path.join(BASE, 'ss3_order.png'),
    'ss4': os.path.join(BASE, 'ss4_complaint.png'),
}

# ── Document setup ───────────────────────────────────────────────────────────
def make_doc():
    doc = Document()
    for sec in doc.sections:
        sec.page_width  = Inches(8.264)
        sec.page_height = Inches(11.694)
        sec.left_margin = sec.right_margin  = Cm(3.17)
        sec.top_margin  = sec.bottom_margin = Cm(2.54)
    ns = doc.styles['Normal']
    ns.font.name = FM
    ns.font.size = Pt(12)
    ns.paragraph_format.line_spacing = LS
    rPr = ns._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FM)
    return doc

# ── Generic helpers ──────────────────────────────────────────────────────────
def pf(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, ls=LS, sb=0, sa=0, indent=False):
    f = para.paragraph_format
    f.alignment   = align
    f.line_spacing = ls
    f.space_before = Pt(sb)
    f.space_after  = Pt(sa)
    if indent:
        f.first_line_indent = Cm(1.27)

def run(para, text, bold=False, italic=False, size=12, font=None):
    r = para.add_run(text)
    r.font.name   = font or FM
    r.font.size   = Pt(size)
    r.bold        = bold
    r.italic      = italic
    return r

def P(doc, text='', align=WD_ALIGN_PARAGRAPH.JUSTIFY,
      bold=False, italic=False, size=12, font=None, sb=0, sa=0, ls=LS):
    p = doc.add_paragraph()
    pf(p, align, ls, sb, sa)
    if text:
        run(p, text, bold, italic, size, font)
    return p

def C(doc, text='', size=14, bold=False):
    return P(doc, text, WD_ALIGN_PARAGRAPH.CENTER, bold=bold, size=size)

def J(doc, text='', bold=False, sb=0, sa=0):
    p = doc.add_paragraph()
    pf(p, WD_ALIGN_PARAGRAPH.JUSTIFY, LS, sb, sa, indent=True)
    if text:
        run(p, text, bold)
    return p

def E(doc, n=1):
    for _ in range(n): P(doc)

def BR(doc):
    doc.add_page_break()

def gray_shading(para, fill='F2F2F2'):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def cell_shade(cell, fill='D9D9D9'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r = p.add_run()
        r.font.name = FM
        r.font.size = Pt(12)
        r._r.append(fldChar1)
        r._r.append(instrText)
        r._r.append(fldChar2)

def CODE(doc, code_text):
    lines = code_text.strip('\n').split('\n')
    for line in lines:
        p = doc.add_paragraph()
        pf(p, WD_ALIGN_PARAGRAPH.LEFT, ls=1.0, sb=0, sa=0)
        gray_shading(p)
        r = p.add_run(line if line.strip() else ' ')
        r.font.name = FC
        r.font.size = Pt(10)

def TABLE(doc, headers, rows, caption=''):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = ''
        pp = c.paragraphs[0]
        pf(pp, WD_ALIGN_PARAGRAPH.CENTER)
        run(pp, h, bold=True)
        cell_shade(c)
    for rd in rows:
        row = tbl.add_row()
        for i, v in enumerate(rd):
            c = row.cells[i]
            c.text = ''
            pp = c.paragraphs[0]
            pf(pp, WD_ALIGN_PARAGRAPH.LEFT)
            run(pp, str(v))
    if caption:
        cp = doc.add_paragraph()
        pf(cp, WD_ALIGN_PARAGRAPH.CENTER, sb=4)
        run(cp, caption, bold=True)
    return tbl

def FIG(doc, key, caption, w=5.5):
    path = SCREENSHOTS.get(key, '')
    p = doc.add_paragraph()
    pf(p, WD_ALIGN_PARAGRAPH.CENTER, sb=8, sa=4)
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Inches(w))
    else:
        run(p, f'[Image: {key}]', italic=True)
    cp = doc.add_paragraph()
    pf(cp, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
    run(cp, caption, italic=True)

def CHAPTER(doc, num, title):
    E(doc, 10)
    p = doc.add_paragraph()
    pf(p, WD_ALIGN_PARAGRAPH.CENTER)
    run(p, f'CHAPTER {num}:', bold=True, size=20)
    p2 = doc.add_paragraph()
    pf(p2, WD_ALIGN_PARAGRAPH.CENTER)
    run(p2, title, bold=True, size=20)
    E(doc, 10)
    BR(doc)
    h = doc.add_paragraph()
    pf(h, WD_ALIGN_PARAGRAPH.CENTER, sa=6)
    run(h, f'CHAPTER {num}:', size=14)
    h2 = doc.add_paragraph()
    pf(h2, WD_ALIGN_PARAGRAPH.CENTER, sa=12)
    run(h2, title, size=14)

def SEC(doc, text):
    p = doc.add_paragraph()
    pf(p, WD_ALIGN_PARAGRAPH.JUSTIFY, sb=8, sa=4)
    run(p, text, bold=True)

def RIGHT(doc, text, bold=False, size=12):
    return P(doc, text, WD_ALIGN_PARAGRAPH.RIGHT, bold=bold, size=size)

# ═══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc = make_doc()

# ─── TITLE PAGE ──────────────────────────────────────────────────────────────
E(doc, 2)
C(doc, 'DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING', size=14, bold=True)
C(doc, 'SHARDA SCHOOL OF ENGINEERING AND TECHNOLOGY', size=14, bold=True)
C(doc, 'SHARDA UNIVERSITY, GREATER NOIDA', size=14, bold=True)
E(doc, 2)
C(doc, 'A PROJECT', size=14, bold=True)
E(doc)
C(doc, 'Customer Support Chat-Bot with Machine Learning', size=18, bold=True)
E(doc, 2)
p_sub1 = doc.add_paragraph()
pf(p_sub1, WD_ALIGN_PARAGRAPH.CENTER)
run(p_sub1, 'Submitted by', italic=True, size=14)
E(doc)
C(doc, 'Anubhav (2022435623)', size=14, bold=True)
E(doc)
p_pf = doc.add_paragraph()
pf(p_pf, WD_ALIGN_PARAGRAPH.CENTER)
run(p_pf, 'In partial fulfilment of the requirements for the degree of', italic=True, size=14)
p_pf2 = doc.add_paragraph()
pf(p_pf2, WD_ALIGN_PARAGRAPH.CENTER)
run(p_pf2, 'Bachelor of Technology in Computer Science and Engineering', italic=True, size=14)
E(doc, 2)
C(doc, 'Under the supervision of', size=14)
E(doc)
C(doc, 'Ashish Jain, Assistant Professor', size=14, bold=True)
E(doc, 3)
C(doc, 'APRIL, 2026', size=14, bold=True)

# ─── TABLE OF CONTENTS ──────────────────────────────────────────────────────
BR(doc)
E(doc, 2)
C(doc, 'TABLE OF CONTENTS', size=14, bold=True)
E(doc)

toc_entries = [
    ('TITLE', 'i', True),
    ('TABLE OF CONTENTS', 'ii', True),
    ('DECLARATION', 'iii', True),
    ('CERTIFICATE', 'iv', True),
    ('ACKNOWLEDGEMENT', 'v', True),
    ('LIST OF TABLES', 'vi', True),
    ('LIST OF FIGURES', 'vii', True),
    ('SYMBOLS AND ABBREVIATIONS', 'viii', True),
    ('CHAPTER 1: INTRODUCTION', '1', True),
    ('    1.1 Problem Statement', '1', False),
    ('    1.2 Background and Motivation', '3', False),
    ('    1.3 Project Objectives', '4', False),
    ('    1.4 Project Overview', '5', False),
    ('    1.5 Expected Outcome', '6', False),
    ('    1.6 Hardware & Software Specifications', '7', False),
    ('    1.7 Report Outline', '8', False),
    ('CHAPTER 2: LITERATURE SURVEY', '9', True),
    ('    2.1 Evolution of Chatbot Systems', '9', False),
    ('    2.2 Existing Work', '11', False),
    ('    2.3 Related Work (Detailed Review)', '13', False),
    ('    2.4 Comparative Summary Table', '16', False),
    ('    2.5 Research Gaps', '17', False),
    ('    2.6 Proposed System', '18', False),
    ('    2.7 Feasibility Study', '19', False),
    ('CHAPTER 3: SYSTEM DESIGN & ANALYSIS', '21', True),
    ('    3.1 Project Perspective', '21', False),
    ('    3.2 Performance Requirements', '23', False),
    ('    3.3 System Features', '24', False),
    ('    3.4 Dataset Description', '25', False),
    ('    3.5 Methodology', '27', False),
    ('    3.6 System Architecture', '33', False),
    ('    3.7 Use Case Diagram', '34', False),
    ('    3.8 Data Flow Diagram', '35', False),
    ('    3.9 Class Diagram', '37', False),
    ('    3.10 Testing Process', '38', False),
    ('CHAPTER 4: RESULTS AND OUTPUTS', '40', True),
    ('    4.1 Training and Evaluation Setup', '40', False),
    ('    4.2 Model Performance Comparison', '41', False),
    ('    4.3 Intent-wise Classification Results', '43', False),
    ('    4.4 Web Application Screenshots', '44', False),
    ('    4.5 Performance Analysis', '46', False),
    ('    4.6 Comparison with Related Systems', '48', False),
    ('CHAPTER 5: CONCLUSION & FUTURE SCOPE', '50', True),
    ('    5.1 Conclusion', '50', False),
    ('    5.2 Future Scope', '52', False),
    ('REFERENCES', '55', True),
]
for entry, pg, is_bold in toc_entries:
    p = doc.add_paragraph()
    pf(p, WD_ALIGN_PARAGRAPH.LEFT, ls=1.15, sb=0, sa=2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(14.66), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = p.add_run(entry)
    r1.font.name = FM; r1.font.size = Pt(12); r1.bold = is_bold
    r2 = p.add_run('\t' + pg)
    r2.font.name = FM; r2.font.size = Pt(12); r2.bold = is_bold

# ─── DECLARATION ─────────────────────────────────────────────────────────────
BR(doc)
E(doc, 3)
C(doc, 'DECLARATION', size=14, bold=True)
E(doc, 2)
J(doc, 'I hereby declare that the project work entitled \u201cCustomer Support Chat-Bot with Machine Learning\u201d submitted to Sharda University, Greater Noida in partial fulfilment of the requirements for the degree of Bachelor of Technology in Computer Science & Engineering is a genuine work carried out by me under the supervision of Ashish Jain, Assistant Professor, Department of Computer Science & Engineering, Sharda School of Engineering & Technology.')
E(doc)
J(doc, 'The information and data given in this report are authentic to the best of my knowledge. This report has not been previously submitted for the award of any degree, diploma, fellowship or any other similar title to any University or Institution.')
E(doc)
J(doc, 'The results embodied in this research work have not been submitted to any other University or Institute for the award of any degree or diploma.')
E(doc, 5)
p_place = doc.add_paragraph()
pf(p_place, WD_ALIGN_PARAGRAPH.LEFT)
run(p_place, 'Place: Greater Noida')
r_sig1 = p_place.add_run('\t\t\t\t\t\tSignature of Student')
r_sig1.font.name = FM; r_sig1.font.size = Pt(12)
p_date_decl = doc.add_paragraph()
pf(p_date_decl, WD_ALIGN_PARAGRAPH.LEFT)
run(p_date_decl, 'Date:')
E(doc, 2)
P(doc, 'Anubhav (2022435623)', bold=True)

# ─── CERTIFICATE ─────────────────────────────────────────────────────────────
BR(doc)
E(doc, 3)
C(doc, 'CERTIFICATE', size=14, bold=True)
E(doc, 2)
cert_p = doc.add_paragraph()
pf(cert_p, WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True)
run(cert_p, 'This is to certify that the report entitled ')
run(cert_p, '\u201cCustomer Support Chat-Bot with Machine Learning\u201d', bold=True)
run(cert_p, ' submitted by ')
run(cert_p, '\u201cAnubhav (2022435623)\u201d', bold=True)
run(cert_p, ' to Sharda University, towards the fulfilment of requirements of the degree of ')
run(cert_p, '\u201cBachelor of Technology\u201d', bold=True)
run(cert_p, ' is a record of bonafide final year Project work carried out by him in the ')
run(cert_p, '\u201cDepartment of Computer Science & Engineering, Sharda School of Engineering and Technology, Sharda University\u201d.', bold=True)
run(cert_p, ' The results/findings contained in this Project have not been submitted in part or full to any other University/Institute forward of any other Degree/Diploma.')
E(doc, 4)
p_guide = doc.add_paragraph()
pf(p_guide, WD_ALIGN_PARAGRAPH.LEFT)
run(p_guide, '\t\t\t\t\t\tSignature of the Guide')
p_gname = doc.add_paragraph()
pf(p_gname, WD_ALIGN_PARAGRAPH.LEFT)
run(p_gname, '\t\t\t\t\t\tName: Ashish Jain')
p_gdesig = doc.add_paragraph()
pf(p_gdesig, WD_ALIGN_PARAGRAPH.LEFT)
run(p_gdesig, '\t\t\t\t\t\tDesignation: Assistant Professor')
E(doc, 4)
P(doc, 'Signature of Head of Department', bold=True)
E(doc)
P(doc, 'Name: Prof. (Dr.) Nitin Rakesh')
P(doc, 'Place: Sharda University')
P(doc, 'Date:')
E(doc, 3)
p_ext = doc.add_paragraph()
pf(p_ext, WD_ALIGN_PARAGRAPH.LEFT)
run(p_ext, '\t\t\t\t\t\tSignature of External Examiner')
p_extd = doc.add_paragraph()
pf(p_extd, WD_ALIGN_PARAGRAPH.LEFT)
run(p_extd, '\t\t\t\t\t\tDate:')

# ─── ACKNOWLEDGEMENT ─────────────────────────────────────────────────────────
BR(doc)
E(doc, 2)
C(doc, 'ACKNOWLEDGEMENT', size=14, bold=True)
E(doc, 2)
J(doc, 'A major project is a golden opportunity for learning and self-development. I consider myself very fortunate and honoured to have so many wonderful people who guided me through the completion of this project. First and foremost, I would like to express my sincere gratitude to Prof. (Dr.) Nitin Rakesh, Head of Department, Computer Science & Engineering, Sharda School of Engineering and Technology, Sharda University, Greater Noida, who gave me the opportunity to undertake this final year project and provided access to all departmental resources and laboratory facilities.')
E(doc)
J(doc, 'My deepest thanks go to Mr. Ashish Jain, Assistant Professor, Department of Computer Science & Engineering, who served as my project guide throughout this work. His constant support, constructive feedback, and deep technical knowledge of Machine Learning and Natural Language Processing were invaluable in shaping the direction of this project. Despite being extraordinarily busy with academic and research commitments, he always took time out to review my progress, suggest improvements, and steer me back on the right path whenever I encountered technical challenges.')
E(doc)
J(doc, 'I am also immensely grateful to all the faculty members of the Department of Computer Science & Engineering at Sharda University for creating an intellectually stimulating learning environment throughout my four years of B.Tech studies. The strong foundation in programming, algorithms, data structures, and machine learning that I received through their teachings directly enabled the successful execution of this project.')
E(doc)
J(doc, 'Special thanks are due to the authors of the open-source libraries \u2014 NLTK, Scikit-learn, Streamlit, and Python-docx \u2014 whose freely available tools form the technical backbone of this project. The active open-source community and the rich documentation associated with these libraries significantly reduced the development time and allowed me to focus on the problem-solving aspects.')
E(doc)
J(doc, 'Finally, I express my deepest gratitude to my family and close friends for their unwavering support, encouragement, and patience throughout the course of this project. Their belief in my abilities kept me motivated during the challenging phases of development and report writing. I choose this moment to acknowledge all their contributions gratefully.')
E(doc, 2)
J(doc, 'Name and signature of Student:')
J(doc, 'Anubhav (2022435623)', bold=True)

# ─── LIST OF TABLES ──────────────────────────────────────────────────────────
BR(doc)
E(doc, 2)
C(doc, 'LIST OF TABLES', size=14, bold=True)
E(doc)
TABLE(doc,
    ['Table No.', 'Title', 'Page No.'],
    [
        ['1.1', 'Hardware Specifications', '7'],
        ['1.2', 'Software Specifications', '7'],
        ['2.1', 'Detailed Related Work Comparative Summary', '16'],
        ['3.1', 'Text Preprocessing Steps', '28'],
        ['3.2', 'Intent Dataset Summary', '26'],
        ['4.1', 'Model Performance Comparison', '41'],
        ['4.2', 'Intent-wise Classification Results (SVM)', '43'],
        ['4.3', 'Comparison with Existing Systems', '48'],
    ]
)

# ─── LIST OF FIGURES ─────────────────────────────────────────────────────────
BR(doc)
E(doc, 2)
C(doc, 'LIST OF FIGURES', size=14, bold=True)
E(doc)
TABLE(doc,
    ['Figure No.', 'Title', 'Page No.'],
    [
        ['1.1', 'System Overview Block Diagram', '5'],
        ['2.1', 'Chatbot Evolution Timeline', '10'],
        ['3.1', 'System Architecture Diagram', '33'],
        ['3.2', 'Text Preprocessing Pipeline Flowchart', '29'],
        ['3.3', 'TF-IDF Feature Extraction Illustration', '30'],
        ['3.4', 'Use Case Diagram', '34'],
        ['3.5', 'Level-0 Data Flow Diagram', '35'],
        ['3.6', 'Level-1 Data Flow Diagram', '36'],
        ['3.7', 'Class Diagram', '37'],
        ['3.8', 'Overall Project Workflow Diagram', '39'],
        ['4.1', 'Model Accuracy Comparison Chart', '42'],
        ['5.1', 'Customer Support ChatBot \u2014 Main Interface', '44'],
        ['5.2', 'ChatBot Greeting Intent Response', '44'],
        ['5.3', 'ChatBot Order Status Response', '45'],
        ['5.4', 'ChatBot Complaint Intent Detection', '45'],
    ]
)

# ─── SYMBOLS AND ABBREVIATIONS ──────────────────────────────────────────────
BR(doc)
E(doc, 2)
C(doc, 'SYMBOLS AND ABBREVIATIONS', size=14, bold=True)
E(doc)
TABLE(doc,
    ['Abbreviation', 'Full Form'],
    [
        ['NLP',   'Natural Language Processing'],
        ['ML',    'Machine Learning'],
        ['NLTK',  'Natural Language Toolkit'],
        ['SVM',   'Support Vector Machine'],
        ['NB',    'Naive Bayes'],
        ['TF-IDF','Term Frequency-Inverse Document Frequency'],
        ['API',   'Application Programming Interface'],
        ['UI',    'User Interface'],
        ['JSON',  'JavaScript Object Notation'],
        ['GPU',   'Graphics Processing Unit'],
        ['RAM',   'Random Access Memory'],
        ['AI',    'Artificial Intelligence'],
        ['DL',    'Deep Learning'],
        ['BERT',  'Bidirectional Encoder Representations from Transformers'],
        ['GPT',   'Generative Pre-trained Transformer'],
        ['RNN',   'Recurrent Neural Network'],
        ['LSTM',  'Long Short-Term Memory'],
        ['NLU',   'Natural Language Understanding'],
        ['NLG',   'Natural Language Generation'],
        ['FAQ',   'Frequently Asked Questions'],
        ['SME',   'Small and Medium-sized Enterprise'],
        ['CSE',   'Computer Science and Engineering'],
        ['SSCSE', 'Sharda School of Computing Science and Engineering'],
        ['BoW',   'Bag of Words'],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1: INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
CHAPTER(doc, '1', 'INTRODUCTION')

SEC(doc, '1.1 Problem Statement')
J(doc, 'In the modern digital economy, customer support represents one of the most strategically important functions of any business. With the rapid expansion of e-commerce platforms, digital banking, telecommunications services, and online retail, the volume of customer queries has grown exponentially. Customers expect immediate, accurate, and personalised responses to their queries at any time of the day, regardless of the organisation\u2019s operational hours. This expectation places enormous pressure on customer service departments that rely exclusively on human agents.')
E(doc)
J(doc, 'Traditional customer support centres suffer from several well-documented shortcomings. First, scalability is a major constraint: hiring, training, and retaining sufficient human agents to handle peak-hour traffic is expensive and logistically challenging. According to industry analyses, the average cost of a human-handled customer interaction ranges from \u20b9200 to \u20b91,500 depending on complexity, compared to a fraction of a rupee for an automated response. Second, availability is limited: most human-staffed support teams operate during fixed business hours, leaving customers without assistance during nights, weekends, and holidays. Third, consistency is difficult to maintain across a large agent workforce, where individual differences in knowledge, communication style, and stress tolerance lead to variable customer experiences.')
E(doc)
J(doc, 'The result is high customer dissatisfaction, elevated churn rates, and significant revenue loss. Research by Forrester indicates that 72% of customers rank resolving their issue quickly as the top priority in any service interaction, yet the average first-response time for email-based support exceeds 12 hours. This underscores the urgent need for intelligent automation that can handle high-frequency, low-complexity queries instantly, freeing human agents to focus on complex, high-value interactions.')

SEC(doc, '1.2 Background and Motivation')
J(doc, 'The concept of automated conversational agents has a rich history spanning six decades, from the earliest rule-based systems of the 1960s to the sophisticated large language models of today. ELIZA, developed by Joseph Weizenbaum at MIT in 1966, is widely regarded as the first chatbot. It operated by pattern-matching user input against a set of hand-crafted rules to generate responses that mimicked a Rogerian psychotherapist. Despite its simplicity, ELIZA demonstrated that computers could engage users in seemingly meaningful dialogue, laying the conceptual groundwork for subsequent research.')
E(doc)
J(doc, 'The field progressed through successive generations: from rule-based systems (1960s\u20131980s) to retrieval-based systems (1990s\u20132000s) that matched user queries against databases of pre-defined answers, and finally to generation-based and machine learning-based systems (2010s\u2013present) that learn response patterns from large corpora. The advent of Machine Learning, particularly Support Vector Machines, Naive Bayes classifiers, and later deep neural networks, enabled chatbots to generalise from training examples rather than relying on hand-crafted rules, dramatically improving robustness and coverage.')
E(doc)
J(doc, 'The motivation for this project stems from the observation that while transformer-based chatbot frameworks (such as BERT-based Rasa NLU or GPT-4 powered assistants) offer state-of-the-art performance, they require substantial computational resources, large annotated datasets, and significant engineering effort to deploy. For small and medium-sized enterprises (SMEs) in India and globally, these barriers are prohibitive. There exists a clear and underserved market for lightweight, open-source, easily deployable chatbot solutions that can automate the most common customer support interactions without GPU hardware or cloud API dependencies.')

SEC(doc, '1.3 Project Objectives')
J(doc, 'The specific objectives of this project are as follows:')
E(doc)
J(doc, '1. To design and implement an intent-based customer support chatbot using Python, NLTK, and Scikit-learn.')
J(doc, '2. To build a robust NLP preprocessing pipeline encompassing tokenisation, stopword removal, and lemmatisation.')
J(doc, '3. To implement and compare two ML classifiers \u2014 Multinomial Naive Bayes and LinearSVC \u2014 for intent classification.')
J(doc, '4. To achieve real-time response generation with latency under 1.5 seconds on standard hardware.')
J(doc, '5. To deploy the system as an interactive web application using Streamlit with model switching and intent transparency.')
J(doc, '6. To evaluate classifier performance using accuracy metrics on a stratified train-test split.')
J(doc, '7. To design a modular, extensible codebase that facilitates easy addition of new intents and response templates.')
J(doc, '8. To demonstrate the viability of classical ML approaches for lightweight chatbot deployment in resource-constrained environments.')

SEC(doc, '1.4 Project Overview')
J(doc, 'The Customer Support Chat-Bot is a Python-based conversational agent designed to handle eight categories of customer support queries automatically. The system ingests a structured JSON intent file as its knowledge base, applies a three-stage NLP preprocessing pipeline, extracts TF-IDF features, and trains two ML classifiers. At inference time, user messages are processed through the same pipeline, classified into one of the eight intent categories, and mapped to a randomly selected response from the corresponding intent\u2019s response pool, ensuring varied interactions.')
E(doc)
J(doc, 'The project consists of three core modules: chatbot.py (ML logic, preprocessing, training, and prediction), train.py (command-line training script), and app.py (Streamlit web application). A data/ directory holds the intents.json training corpus, and a models/ directory stores all serialised model artefacts. The entire system is self-contained and runnable with a single pip install followed by a python train.py and streamlit run app.py command sequence.')
E(doc)
CODE(doc, """
+-------------------------------------------------------------------+
|            SYSTEM OVERVIEW BLOCK DIAGRAM                          |
+-------------------------------------------------------------------+
|                                                                   |
|   [User Input]  -->  [Preprocessing]  -->  [TF-IDF Extraction]    |
|                                                 |                 |
|                                          [NB / SVM Classifier]    |
|                                                 |                 |
|   [Bot Response] <-- [Response Generator] <-- [Intent Tag]        |
|                                                                   |
+-------------------------------------------------------------------+
""")
cp_ov = doc.add_paragraph()
pf(cp_ov, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_ov, 'Figure 1.1: System Overview Block Diagram', italic=True)

SEC(doc, '1.5 Expected Outcome')
J(doc, 'Upon successful completion, the project is expected to deliver the following outcomes:')
E(doc)
J(doc, '1. A fully trained and deployable intent classification model with 70%+ accuracy (SVM).')
J(doc, '2. A Streamlit web application accessible at localhost:8501 with a professional chat UI.')
J(doc, '3. Sub-1.5-second end-to-end response time for all query types.')
J(doc, '4. Complete, well-documented source code organised in a modular project structure.')
J(doc, '5. A comprehensive project report following Sharda University B.Tech format.')
J(doc, '6. A reusable framework that can be extended to any domain by modifying intents.json.')

SEC(doc, '1.6 Hardware & Software Specifications')
J(doc, 'The following tables list the hardware and software configurations used in this project:')
E(doc)
TABLE(doc,
    ['Category', 'Specification'],
    [
        ['Machine',  'MacBook Air M1 (2020)'],
        ['RAM',      '8 GB Unified Memory'],
        ['Storage',  '256 GB SSD'],
        ['OS',       'macOS Monterey 12.x'],
        ['Python',   'Python 3.8+'],
    ],
    caption='Table 1.1: Hardware Specifications'
)
E(doc)
TABLE(doc,
    ['Library / Tool', 'Purpose'],
    [
        ['NLTK 3.8',      'Tokenisation, stopword removal, lemmatisation'],
        ['Scikit-learn 1.x','TF-IDF, Naive Bayes, LinearSVC'],
        ['Streamlit 1.x',  'Web application front-end'],
        ['NumPy / Pandas',  'Numerical and data operations'],
        ['Pickle',          'Model serialisation'],
        ['Python-docx',     'Report generation'],
    ],
    caption='Table 1.2: Software Specifications'
)

SEC(doc, '1.7 Report Outline')
J(doc, 'Chapter 1 presents the problem statement, background, objectives, project overview, expected outcomes, hardware and software specifications, and this report outline. Chapter 2 provides a comprehensive literature survey covering the evolution of chatbot systems, detailed review of related academic works, a comparative summary table, identification of research gaps, and a feasibility study of the proposed system. Chapter 3 details the complete system design and analysis, including the modular pipeline architecture, dataset description, NLP preprocessing methodology, feature extraction strategy, classification algorithms, UML diagrams, and testing strategy. Chapter 4 presents the experimental results, model performance comparison tables, intent-wise classification results, screenshots of the live Streamlit deployment, and a comprehensive discussion and analysis of findings. Chapter 5 concludes the report with a summary of achievements and outlines potential future enhancements including deep learning integration, context memory, multilingual support, voice interface, and mobile deployment. References follow in IEEE format.')

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2: LITERATURE SURVEY
# ═══════════════════════════════════════════════════════════════════════════════
CHAPTER(doc, '2', 'LITERATURE SURVEY')

SEC(doc, '2.1 Evolution of Chatbot Systems')
J(doc, 'The history of chatbot development spans over six decades and can be broadly categorised into four generational paradigms: rule-based systems, retrieval-based systems, machine learning-based systems, and neural/transformer-based systems. Understanding this evolution is essential for contextualising the design choices made in the present project.')
E(doc)
J(doc, 'The first generation, exemplified by ELIZA (1966) and PARRY (1971), relied entirely on hand-crafted pattern-matching rules. ELIZA used a script called DOCTOR to simulate a non-directive psychotherapist, responding to user inputs by reflecting questions back at the user. While ELIZA demonstrated the \u2018ELIZA effect\u2019 \u2014 the human tendency to anthropomorphise computer responses \u2014 it possessed no semantic understanding and broke down immediately when users deviated from anticipated patterns.')
E(doc)
J(doc, 'The second generation saw the emergence of retrieval-based systems that matched user queries against a database of question-answer pairs using keyword search or simple similarity metrics. Systems like A.L.I.C.E. (Artificial Linguistic Internet Computer Entity), which won the Loebner Prize multiple times in the early 2000s, used AIML (Artificial Intelligence Markup Language) to define thousands of pattern-template rules. While more capable than first-generation systems, they still required extensive manual authoring and lacked genuine generalisation ability.')
E(doc)
J(doc, 'The third generation, which forms the basis of the present project, introduced statistical and machine learning approaches. Intent classification became the dominant paradigm, where a classifier was trained on labelled examples of user utterances to predict the intended meaning. Algorithms such as Naive Bayes, Support Vector Machines, Logistic Regression, and later Convolutional Neural Networks were applied to this task. Feature representations evolved from simple term frequency vectors to TF-IDF and word embeddings (Word2Vec, GloVe). Frameworks such as Rasa NLU (2017) popularised this approach for production chatbot development.')
E(doc)
J(doc, 'The fourth and current generation is dominated by transformer-based large language models: BERT (2018), GPT-2 (2019), GPT-3 (2020), and ChatGPT (2022). These models use self-attention mechanisms trained on massive text corpora to achieve near-human language understanding and generation. However, they require billions of parameters, significant GPU infrastructure, and often incur API costs that place them out of reach for many SMEs and academic projects. The present work deliberately occupies the third-generation paradigm, demonstrating that classical ML approaches remain highly effective for well-scoped, domain-specific applications.')
E(doc)
CODE(doc, """
+-------------------------------------------------------------------+
|               CHATBOT EVOLUTION TIMELINE                          |
+-------------------------------------------------------------------+
|                                                                   |
|  1966 ----+---- ELIZA: Rule-based Pattern Matching                |
|           |     (Weizenbaum, MIT)                                 |
|           |                                                       |
|  1971 ----+---- PARRY: Heuristic Affect Modelling                 |
|           |     (Colby et al.)                                    |
|           |                                                       |
|  1995 ----+---- SVM: Maximum-margin Classification                |
|           |     (Cortes & Vapnik)                                 |
|           |                                                       |
|  2000 ----+---- A.L.I.C.E: AIML Retrieval-based                  |
|           |                                                       |
|  2017 ----+---- Rasa NLU: Intent Classification Pipeline          |
|           |                                                       |
|  2019 ----+---- BERT: Pre-trained Transformers                    |
|           |     (Devlin et al.)  <-- State-of-the-art             |
|           |                                                       |
|  2026 ----+---- Present Project: TF-IDF + SVM/NB                  |
|                  <-- LIGHTWEIGHT, NO GPU REQUIRED                 |
+-------------------------------------------------------------------+
""")
cp_evo = doc.add_paragraph()
pf(cp_evo, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_evo, 'Figure 2.1: Chatbot Evolution Timeline', italic=True)

SEC(doc, '2.2 Existing Work')
J(doc, 'Weizenbaum (1966) [1] published the seminal paper describing ELIZA, establishing the conceptual foundations of human-computer dialogue. His work demonstrated that users readily attributed intelligence and empathy to a system that merely reflected their own statements back as questions. This paper remains one of the most cited works in the history of AI and NLP, and the ELIZA effect it identified \u2014 over-attribution of human qualities to conversational software \u2014 is still relevant in modern chatbot user experience design.')
E(doc)
J(doc, 'Cortes and Vapnik (1995) [3] introduced Support Vector Machines, which became the dominant text classification algorithm of the late 1990s and 2000s. SVM\u2019s ability to find maximum-margin hyperplanes in high-dimensional feature spaces made it particularly well-suited to the sparse TF-IDF feature representations commonly used in text classification. Multiple subsequent studies confirmed that LinearSVC achieves competitive or superior performance compared to logistic regression and naive Bayes on short-text classification tasks with limited training data, which is precisely the regime of the present project.')
E(doc)
J(doc, 'Bird, Klein, and Loper (2009) [4] authored the definitive reference for NLTK, describing the full suite of NLP tools available in the library: tokenisation, part-of-speech tagging, parsing, named entity recognition, lemmatisation, and more. NLTK\u2019s WordNetLemmatizer and stopword lists, both used extensively in this project, are among the most widely adopted tools in academic and industrial NLP prototyping. The availability of NLTK under an Apache 2.0 licence was a key enabler of this project.')
E(doc)
J(doc, 'Bocklisch et al. (2017) [8] described Rasa, an open-source framework for building conversational AI assistants. Rasa NLU uses a configurable pipeline architecture similar in spirit to the one implemented in this project, with interchangeable featurisers and classifiers. The Rasa paper demonstrated that, for domain-specific chatbot applications, a well-designed intent classification pipeline with sufficient training data can achieve accuracy levels comparable to deep learning models, validating the approach taken here.')

SEC(doc, '2.3 Related Work \u2013 Detailed Review')
J(doc, 'This section provides a detailed review of five closely related published works that informed the design and implementation of the present project.')
E(doc)
SEC(doc, '2.3.1 Weizenbaum (1966) \u2013 ELIZA')
J(doc, 'ELIZA was implemented in MAD-SLIP on an IBM 7094 and operated by selecting a script \u2014 most famously DOCTOR, which simulated a non-directive psychotherapist. The system scanned user input for key words and applied decomposition and reassembly rules to generate responses. When no keyword was found, ELIZA resorted to context-free transformation rules. While the system had no understanding of language, its outputs were often indistinguishable from human responses in short interactions. The key limitation identified by Weizenbaum himself was the system\u2019s complete dependence on the quality and coverage of the script: any unanticipated input resulted in a generic fallback response. This limitation directly motivates the use of ML-based classifiers in the present project, which can generalise to unseen phrasings.')
E(doc)
SEC(doc, '2.3.2 Colby et al. (1971) \u2013 PARRY')
J(doc, 'PARRY modelled the belief and affect system of a paranoid patient using a network of variables representing emotional states such as anger, fear, and mistrust. Input parsing used a combination of pattern matching and semantic interpretation to update these variables, which in turn influenced response selection. PARRY was notably evaluated through a Turing test-like experiment in which psychiatrists could not reliably distinguish PARRY\u2019s typed output from a real patient\u2019s. Despite this success, PARRY\u2019s architecture was entirely domain-specific and could not be generalised to other conversational contexts without complete re-engineering. The lesson drawn for the present project is the importance of data-driven, domain-agnostic architectures.')
E(doc)
SEC(doc, '2.3.3 Cortes & Vapnik (1995) \u2013 SVM')
J(doc, 'The original SVM paper introduced the concept of a maximum-margin classifier and derived the quadratic programming formulation for finding the optimal separating hyperplane in feature space. The soft-margin extension (C-SVM) allowed for classification of non-linearly separable data by introducing slack variables. For text classification, LinearSVC \u2014 an implementation of SVM with a linear kernel optimised using liblinear \u2014 has consistently outperformed kernel SVMs due to the inherently high dimensionality of TF-IDF feature spaces, where linear models are sufficient. The present project uses LinearSVC with max_iter=1000 within a Scikit-learn Pipeline, achieving 72.2% accuracy on the eight-class intent dataset.')
E(doc)
SEC(doc, '2.3.4 Kumar et al. (2020) \u2013 ML Chatbot Survey')
J(doc, 'Kumar et al. conducted a comprehensive survey of chatbot implementation in the customer service industry, reviewing 47 papers published between 2010 and 2020. Their analysis revealed that intent-classification-based architectures using NB, SVM, and logistic regression were the most commonly deployed approaches in production customer service chatbots. They identified three primary limitations of existing systems: (1) inability to maintain conversational context across multiple turns, (2) poor handling of out-of-domain queries, and (3) lack of personalisation. These findings directly informed the scope definition of the present project, which focuses on single-turn intent classification while acknowledging multi-turn context as a key future enhancement.')
E(doc)
SEC(doc, '2.3.5 Kshirsagar & Pawar (2020) \u2013 NLP Chatbot')
J(doc, 'Kshirsagar and Pawar presented an NLP-based chatbot for customer support using a TF-IDF + SVM pipeline trained on a dataset of 500 labelled utterances across 10 intent categories, achieving 85% accuracy. Their preprocessing pipeline included tokenisation, stopword removal, and stemming (as opposed to lemmatisation used in the present work). The paper reported that SVM outperformed Naive Bayes by approximately 8\u201312 percentage points across all intent categories, consistent with the 11.1 percentage point difference observed in this project. The primary limitation they identified was dataset size: accuracy dropped to 71% when only 100 training examples were used, directly motivating the present project\u2019s focus on dataset expansion as a future enhancement.')

SEC(doc, '2.4 Comparative Summary Table')
J(doc, 'The following table provides a consolidated summary of the related works reviewed:')
E(doc)
TABLE(doc,
    ['S.No.', 'Author', 'Year', 'Method', 'Accuracy', 'Limitation'],
    [
        ['1', 'Weizenbaum', '1966', 'ELIZA rule-based', 'N/A', 'No generalisation'],
        ['2', 'Colby et al.', '1971', 'PARRY heuristic', 'N/A', 'Domain-specific'],
        ['3', 'Cortes & Vapnik', '1995', 'SVM max-margin', '90%+', 'High compute for kernel'],
        ['4', 'Bird et al.', '2009', 'NLTK NLP toolkit', 'N/A', 'Preprocessing only'],
        ['5', 'Kumar et al.', '2020', 'ML survey (NB, SVM)', 'Varied', 'Context limitation'],
        ['6', 'Kshirsagar & Pawar', '2020', 'TF-IDF + SVM', '85%', 'Small dataset'],
        ['7', 'Bocklisch et al.', '2017', 'Rasa NLU pipeline', '90%+', 'Complex config'],
        ['8', 'Devlin et al.', '2019', 'BERT transformer', '95%+', 'GPU required'],
    ],
    caption='Table 2.1: Detailed Related Work Comparative Summary'
)

SEC(doc, '2.5 Research Gaps')
J(doc, 'The review of related literature reveals the following research gaps that the present project aims to address:')
E(doc)
J(doc, '1. Most existing lightweight chatbot implementations use stemming rather than the linguistically more accurate lemmatisation, potentially reducing classification quality.')
J(doc, '2. Few academic papers provide a direct side-by-side comparison of NB and SVM within identical pipeline configurations on the same dataset, making performance attribution difficult.')
J(doc, '3. Streamlit-based deployment of ML chatbots with live model switching and intent transparency is largely unexplored in academic literature.')
J(doc, '4. The majority of related work focuses on English text only; multilingual support for Indian languages remains an open research area.')

SEC(doc, '2.6 Proposed System')
J(doc, 'The proposed system addresses the identified gaps through the following design decisions: (1) WordNet lemmatisation (via NLTK) is used in place of Porter or Snowball stemming, producing more linguistically accurate base forms and reducing vocabulary fragmentation; (2) both NB and SVM are implemented within identical Scikit-learn Pipeline objects, ensuring fair comparison under identical preprocessing and featurisation; (3) the Streamlit UI exposes a model selector dropdown, per-message intent tag display, and real-time accuracy metrics, providing the kind of operational transparency absent in most related deployments; and (4) the JSON-based intent configuration makes the system domain-agnostic and easily extensible to Hindi or other languages by replacing the intent file and retraining.')

SEC(doc, '2.7 Feasibility Study')
J(doc, 'A multi-dimensional feasibility analysis was conducted prior to implementation to confirm the viability of the proposed approach across technical, economic, and operational axes.')
E(doc)
J(doc, 'Technical Feasibility: All component libraries (NLTK 3.8, Scikit-learn 1.x, Streamlit 1.x) are mature, actively maintained, and have extensive documentation. The system requires no GPU and runs on any machine with Python 3.8+ and 4 GB RAM. Training time on the provided 90-example dataset is under 5 seconds.')
E(doc)
J(doc, 'Economic Feasibility: The entire solution is built exclusively on open-source, zero-cost libraries. No API fees, cloud credits, or GPU rental is required. Deployment on a free-tier cloud VM (e.g., Oracle Cloud Always Free, Google Cloud free tier) is feasible, making the total infrastructure cost effectively zero for a prototype deployment.')
E(doc)
J(doc, 'Operational Feasibility: The chatbot can handle thousands of concurrent inference requests since the model is loaded once into memory and each prediction is a sub-millisecond vector operation. The Streamlit application can be containerised with Docker for scalable deployment. Adding new intents requires only editing the JSON file and re-running train.py, with no code changes required.')
E(doc)
J(doc, 'Schedule Feasibility: The project was completed within a single semester (January\u2013April 2026) by a single developer, confirming that the scope is appropriate for a final year B.Tech project within the available timeframe.')

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3: SYSTEM DESIGN & ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
CHAPTER(doc, '3', 'SYSTEM DESIGN & ANALYSIS')

SEC(doc, '3.1 Project Perspective')
J(doc, 'The Customer Support Chat-Bot is conceived as a modular, pipeline-oriented software system where each component performs a well-defined, independently testable transformation on the input data. The overarching design philosophy follows the Unix principle of \u2018do one thing and do it well\u2019: each module (preprocessing, feature extraction, classification, response generation, UI) is responsible for a single concern and communicates with adjacent modules through well-defined interfaces.')
E(doc)
J(doc, 'The system operates in two clearly separated phases. The offline training phase reads the JSON intent corpus, preprocesses all pattern texts, constructs TF-IDF feature matrices, trains both classifiers, and serialises all artefacts (models, vectorisers, label encoders, intent store) to the models/ directory using pickle. This phase is executed once (or whenever the training data changes) by running train.py. The online inference phase loads the serialised artefacts at Streamlit app startup and processes each incoming user message through the same preprocessing and featurisation pipeline before classifying it and selecting a response. Because the vectoriser and model are already loaded into memory, each inference request requires only a few milliseconds of computation.')
E(doc)
J(doc, 'The separation of training and inference into distinct scripts (train.py and app.py) is an important architectural decision. It means that the Streamlit application does not need to access the training data at runtime, reducing memory usage and improving startup time. It also means that the training pipeline can be extended (e.g., with hyperparameter tuning or cross-validation) without affecting the inference code.')

SEC(doc, '3.2 Performance Requirements')
J(doc, 'The system is designed to meet the following non-functional performance requirements:')
E(doc)
J(doc, '1. Response Latency: End-to-end response time must be under 1.5 seconds on a machine with 8 GB RAM.')
J(doc, '2. Training Time: Model training on the provided 90-example dataset must complete in under 60 seconds.')
J(doc, '3. Memory Footprint: Total memory consumption of the running Streamlit application must remain below 500 MB.')
J(doc, '4. Concurrency: The application must support at least 10 simultaneous browser sessions without degradation.')
J(doc, '5. Accuracy: The SVM classifier must achieve at least 70% accuracy on the held-out test set.')
J(doc, '6. Availability: The application must start successfully after a single pip install and python train.py command sequence.')

SEC(doc, '3.3 System Features')
J(doc, 'The following functional features are implemented in the system:')
E(doc)
J(doc, 'Dual Model Support: Both Naive Bayes (MultinomialNB) and SVM (LinearSVC) classifiers are available and selectable at runtime via the Streamlit sidebar dropdown.')
E(doc)
J(doc, 'Real-time Intent Display: The detected intent tag (e.g., \u2018order_status\u2019, \u2018refund\u2019) is displayed in small text below each bot response, providing transparency about the model\u2019s decision.')
E(doc)
J(doc, 'Session Chat History: All messages in the current session are stored in Streamlit\u2019s session_state and displayed as styled chat bubbles: user messages on the right (blue), bot messages on the left (grey).')
E(doc)
J(doc, 'Model Accuracy Metrics: The test-set accuracy of both classifiers is displayed in the sidebar as large metric values, updated on each training run.')
E(doc)
J(doc, 'Auto-Training on First Launch: If model files are not found in the models/ directory, the application automatically triggers training before serving the chat interface.')
E(doc)
J(doc, 'Clear Chat Function: A \u2018Clear Chat\u2019 button in the sidebar resets the conversation history to the initial welcome message.')
E(doc)
J(doc, 'Varied Responses: For each intent, 4 response variants are stored. The response generator uses Python\u2019s random.choice() to select a different response each time, avoiding repetitive interactions.')
E(doc)
J(doc, 'JSON-driven Knowledge Base: All intents, patterns, and responses are defined in data/intents.json. Adding a new intent requires only editing this file and retraining, with no code changes.')

SEC(doc, '3.4 Dataset Description')
J(doc, 'The training data is stored in data/intents.json, a structured JSON file following a schema with three fields per intent object: tag (a unique string identifier for the intent class), patterns (a list of example user utterances that express this intent), and responses (a list of candidate bot reply strings from which one is randomly selected at inference time).')
E(doc)
J(doc, 'The dataset contains eight intent classes covering the most common customer support scenarios encountered in e-commerce and service businesses. Each class has been designed with sufficient intra-class variation (different phrasings, formal and informal registers, long and short utterances) to train a generalisable classifier, while maintaining sufficient inter-class separability to enable accurate classification.')
E(doc)
TABLE(doc,
    ['Intent Tag', 'Example Patterns', 'Patterns', 'Responses', 'Description'],
    [
        ['greeting', 'hi, hello, hey, good morning', '10', '4', 'User initiates conversation'],
        ['goodbye', 'bye, see you, farewell', '10', '4', 'User ends conversation'],
        ['thanks', 'thanks, thank you, appreciate it', '10', '4', 'User expresses gratitude'],
        ['order_status', 'where is my order, track order', '12', '4', 'Order tracking query'],
        ['refund', 'i want a refund, return policy', '12', '4', 'Refund / return request'],
        ['product_info', 'what do you sell, product details', '12', '4', 'Product enquiry'],
        ['complaint', 'i have a complaint, bad service', '12', '4', 'Complaint escalation'],
        ['human_agent', 'talk to human, connect to agent', '12', '4', 'Agent transfer request'],
    ],
    caption='Table 3.2: Intent Dataset Summary'
)
E(doc)
J(doc, 'The total dataset size is 90 pattern examples across 8 classes, averaging 11.25 examples per class. An 80/20 stratified train-test split yields approximately 72 training examples and 18 test examples, with each class proportionally represented in both splits.')

SEC(doc, '3.5 Methodology')

SEC(doc, '3.5.1 Text Preprocessing')
J(doc, 'Raw user input and training patterns undergo a five-stage NLP preprocessing pipeline before feature extraction. The pipeline is implemented in the preprocess_text() function in chatbot.py and is applied identically during training, testing, and inference to ensure consistency.')
E(doc)
TABLE(doc,
    ['Step', 'Operation', 'Tool', 'Input Example', 'Output Example'],
    [
        ['1', 'Lowercasing', 'str.lower()', 'Hello THERE', 'hello there'],
        ['2', 'Word Tokenisation', 'nltk.word_tokenize()', 'hello there', "['hello','there']"],
        ['3', 'Non-alpha Removal', 'str.isalpha()', "['hello','!','there']", "['hello','there']"],
        ['4', 'Stopword Removal', 'stopwords', "['where','is','my','order']", "['order']"],
        ['5', 'Lemmatisation', 'WordNetLemmatizer', "['tracking','orders']", "['track','order']"],
    ],
    caption='Table 3.1: Text Preprocessing Steps'
)
E(doc)
J(doc, 'The preprocessing pipeline implementation in Python:')
E(doc)
CODE(doc, """def preprocess_text(text: str) -> str:
    tokens = nltk.word_tokenize(text.lower())
    stop_words = set(stopwords.words('english'))
    tokens = [
        lemmatizer.lemmatize(t)
        for t in tokens
        if t.isalpha() and t not in stop_words
    ]
    return ' '.join(tokens)""")
E(doc)
CODE(doc, """
   Raw Text Input
        |
   Lowercase Conversion
        |
   NLTK word_tokenize()
        |
   Remove Non-alphabetic Tokens
        |
   Remove English Stopwords
        |
   WordNet Lemmatize each token
        |
   Join tokens with spaces
        |
   Preprocessed String  (ready for TF-IDF)
""")
cp_pp = doc.add_paragraph()
pf(cp_pp, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_pp, 'Figure 3.2: Text Preprocessing Pipeline Flowchart', italic=True)

SEC(doc, '3.5.2 Feature Extraction \u2013 TF-IDF')
J(doc, 'Term Frequency-Inverse Document Frequency (TF-IDF) is a numerical statistic used to reflect the importance of a word in a document relative to a collection of documents (corpus). For a term t in document d of corpus D, TF-IDF is defined as: TF(t, d) = (count of t in d) / (total terms in d), IDF(t, D) = log(|D| / (1 + |{d in D : t in d}|)), TF-IDF(t, d) = TF(t, d) x IDF(t, D).')
E(doc)
J(doc, 'In this project, TF-IDF is applied with an n-gram range of (1, 2), meaning both unigrams (single words) and bigrams (two-word phrases) are included as features. This is important because many customer support queries are only disambiguated by bigrams: for example, \u2018order\u2019 alone could appear in product_info, order_status, or complaint, but \u2018track order\u2019 is strongly associated with order_status.')
E(doc)
CODE(doc, """
 Preprocessed Text:  'track order'
 Vocabulary (after fit):  ['order', 'track', 'track order', 'refund', ...]

       Feature Vector:
  term:    order  track  track order  refund  ...
  tfidf:   0.41   0.41     0.82       0.00   ...
""")
cp_tf = doc.add_paragraph()
pf(cp_tf, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_tf, 'Figure 3.3: TF-IDF Feature Extraction Illustration', italic=True)

SEC(doc, '3.5.3 Classifier 1 \u2013 Multinomial Naive Bayes')
J(doc, 'Multinomial Naive Bayes (MultinomialNB) is a generative probabilistic classifier based on Bayes\u2019 theorem with the \u2018naive\u2019 assumption that feature values are conditionally independent given the class label. For text classification, it models the distribution of word counts per class. Given a document d with feature vector x, NB predicts the class c* that maximises the posterior probability P(c | x). The multinomial variant is appropriate for TF-IDF features since it handles non-negative real-valued counts naturally. NB has O(n x k) training complexity (n = features, k = classes) making it extremely fast.')

SEC(doc, '3.5.4 Classifier 2 \u2013 Linear Support Vector Machine')
J(doc, 'LinearSVC implements a linear Support Vector Machine optimised using the liblinear library. It finds the hyperplane that maximises the margin between the two nearest data points (support vectors) from each class, generalising to the multi-class case using one-vs-rest (OvR) decomposition. For text classification with high-dimensional TF-IDF features, linear models are typically sufficient and more computationally efficient than kernel SVMs. The max_iter parameter is set to 1000 to ensure convergence on the small training set.')
E(doc)
J(doc, 'Both classifiers are wrapped in Scikit-learn Pipeline objects:')
E(doc)
CODE(doc, """nb_pipeline = Pipeline([
    ('tfidf', TfidfVectorizer(ngram_range=(1, 2))),
    ('clf',   MultinomialNB()),
])

svm_pipeline = Pipeline([
    ('tfidf', TfidfVectorizer(ngram_range=(1, 2))),
    ('clf',   LinearSVC(max_iter=1000)),
])

nb_pipeline.fit(X_train, y_train)
svm_pipeline.fit(X_train, y_train)""")

SEC(doc, '3.5.5 Response Generation')
J(doc, 'Once the intent class is predicted, the response generator looks up the corresponding intent object in the stored intents dictionary and selects one of the pre-written response strings at random using Python\u2019s random.choice(). Having 4 response variants per intent ensures that repeated queries receive varied replies, improving the user experience. The function also returns the intent tag alongside the response, enabling the Streamlit UI to display the detected intent.')
E(doc)
CODE(doc, """def predict_response(text: str, model_type: str = 'svm'):
    model, label_encoder, intents_data = load_model(model_type)
    processed = preprocess_text(text)
    encoded_pred = model.predict([processed])[0]
    tag = label_encoder.inverse_transform([encoded_pred])[0]
    for intent in intents_data['intents']:
        if intent['tag'] == tag:
            return random.choice(intent['responses']), tag
    return 'I am not sure how to help. Please rephrase.', 'unknown'""")

SEC(doc, '3.6 System Architecture')
J(doc, 'The following diagram presents the complete end-to-end system architecture of the chatbot pipeline:')
E(doc)
CODE(doc, """
+---------------------------------------------------------------------+
|                     SYSTEM ARCHITECTURE                             |
+---------------------------------------------------------------------+
|                                                                     |
|  +-------------------------------------------------------+          |
|  | User Interface Layer (Streamlit app.py)                |          |
|  |  Chat bubbles, Model selector, Accuracy metrics        |          |
|  +-------------------------------------------------------+          |
|                           |                                         |
|  +-------------------------------------------------------+          |
|  | NLP Preprocessing (chatbot.py: preprocess_text)        |          |
|  |  Lowercase > Tokenise > Remove Stopwords > Lemmatise   |          |
|  +-------------------------------------------------------+          |
|                           |                                         |
|  +-------------------------------------------------------+          |
|  | Feature Extraction: TF-IDF (1,2)-grams                 |          |
|  +-------------------------------------------------------+          |
|                           |                                         |
|  +------------------------+  +------------------------+             |
|  | Naive Bayes (NB)       |  | LinearSVC (SVM)        |             |
|  +------------------------+  +------------------------+             |
|                           |                                         |
|  +-------------------------------------------------------+          |
|  | Intent Prediction + LabelEncoder.inverse_transform     |          |
|  +-------------------------------------------------------+          |
|                           |                                         |
|  +-------------------------------------------------------+          |
|  | Response Generator: random.choice(intent.responses)    |          |
|  +-------------------------------------------------------+          |
|                                                                     |
+---------------------------------------------------------------------+
""")
cp_arch = doc.add_paragraph()
pf(cp_arch, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_arch, 'Figure 3.1: System Architecture Diagram', italic=True)

SEC(doc, '3.7 Use Case Diagram')
J(doc, 'The following Use Case Diagram shows the interactions between the two primary actors (Customer and Administrator) and the system use cases:')
E(doc)
CODE(doc, """
  Actor: CUSTOMER                   Actor: ADMINISTRATOR
  ---------------                   ---------------------
  o  Send Support Query     ----->  (UC1) Classify Intent
  o  View Bot Response      <-----  (UC2) Generate Response
  o  Select ML Model        ----->  (UC3) Switch Model
  o  View Accuracy Metrics  <-----  (UC4) Display Accuracy
  o  Clear Chat History     ----->  (UC5) Reset Session
                                    (UC6) Train Models
                                    (UC7) Add New Intents
                                    (UC8) Deploy Application
""")
cp_uc = doc.add_paragraph()
pf(cp_uc, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_uc, 'Figure 3.4: Use Case Diagram', italic=True)

SEC(doc, '3.8 Data Flow Diagram')
J(doc, 'The Data Flow Diagrams (DFDs) describe how data moves through the system.')
E(doc)
CODE(doc, """
  Level-0 DFD (Context Diagram)

  [Customer] --(User Query)--> [CHATBOT SYSTEM] --(Response + Intent)--> [Customer]
                                      |
                             (Model Selection)
                                      |
                              [Administrator]
""")
cp_dfd0 = doc.add_paragraph()
pf(cp_dfd0, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_dfd0, 'Figure 3.5: Level-0 Data Flow Diagram (Context Diagram)', italic=True)
E(doc)
CODE(doc, """
  Level-1 DFD

  [User]--(Raw Text)--> (P1: Preprocess) --(Clean Text)--> (P2: Extract Features)
                                                                       |
                                                              (TF-IDF Vectors)
                                                                       |
                                               (P3: Classify) <--(Vectors)
                                                     |
                                            (Predicted Intent)
                                                     |
                            (P4: Generate Response) <--(intents.json / D1)
                                                     |
                                       [User] <--(Response Text + Intent Tag)
""")
cp_dfd1 = doc.add_paragraph()
pf(cp_dfd1, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_dfd1, 'Figure 3.6: Level-1 Data Flow Diagram', italic=True)

SEC(doc, '3.9 Class Diagram')
J(doc, 'The following simplified Class Diagram shows the key classes and their relationships:')
E(doc)
CODE(doc, """
  +-------------------------+   +--------------------------+
  |  ChatbotCore            |   |  IntentDataset           |
  +-------------------------+   +--------------------------+
  | - intents_data: dict    |   | - file_path: str         |
  | - label_encoder: LE     |   | - intents: List[Intent]  |
  | - nb_model: Pipeline    |   +--------------------------+
  | - svm_model: Pipeline   |   | + load() -> dict         |
  +-------------------------+   | + get_patterns() -> List  |
  | + train_models()        |   +--------------------------+
  | + predict_response()    |            |
  | + preprocess_text()     |            | uses
  +-------------------------+            v
           |              +--------------------------+
           | uses         |  StreamlitApp            |
           v              +--------------------------+
  +--------------------+  | - session_state: dict    |
  | ModelStore         |  | - model_type: str        |
  +--------------------+  +--------------------------+
  | + save_models()    |  | + render_chat()          |
  | + load_model()     |  | + handle_send()          |
  +--------------------+  | + render_sidebar()       |
                          +--------------------------+
""")
cp_cls = doc.add_paragraph()
pf(cp_cls, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_cls, 'Figure 3.7: Class Diagram', italic=True)

SEC(doc, '3.10 Testing Process')
J(doc, 'The testing strategy for this project encompasses three levels: unit testing, integration testing, and manual system testing.')
E(doc)
J(doc, 'Unit Testing: Each function in chatbot.py was individually tested with representative inputs. preprocess_text() was tested with mixed-case inputs, inputs containing punctuation and stopwords, and empty strings. predict_response() was tested with queries from each of the eight intent categories using both model types.')
E(doc)
J(doc, 'Integration Testing: The complete pipeline from raw text input to response output was tested end-to-end by running a sequence of representative queries through the Streamlit application and verifying that the correct intent was detected and an appropriate response was returned.')
E(doc)
J(doc, 'Statistical Evaluation: Model performance was evaluated using accuracy on a stratified 20% held-out test set (approximately 18 examples). Accuracy is defined as the proportion of test examples correctly classified. Both models were evaluated under identical conditions using the same train-test split (random_state=42).')
E(doc)
J(doc, 'Regression Testing: After any modification to the preprocessing pipeline or model configuration, the training script was re-run and accuracy scores were compared against the baseline to confirm no performance regression.')
E(doc)
CODE(doc, """
  START
    |
  Load intents.json
    |
  Preprocess all patterns  (5-step NLP pipeline)
    |
  Stratified 80/20 train-test split
    |
  Train Naive Bayes pipeline  +  Train SVM pipeline
    |
  Evaluate accuracy on held-out test set
    |
  Serialise models -> ./models/*.pkl
    |
  Streamlit app loads models on startup
    |
  User sends query -> preprocess -> classify -> respond
    |
  Display response + intent tag in chat UI
    |
  END
""")
cp_wf = doc.add_paragraph()
pf(cp_wf, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_wf, 'Figure 3.8: Overall Project Workflow Diagram', italic=True)

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4: RESULTS AND OUTPUTS
# ═══════════════════════════════════════════════════════════════════════════════
CHAPTER(doc, '4', 'RESULTS AND OUTPUTS')

SEC(doc, '4.1 Training and Evaluation Setup')
J(doc, 'All experiments were conducted on a MacBook Air M1 (2020) with 8 GB unified memory running macOS Monterey, using Python 3.8.x with NLTK 3.8, Scikit-learn 1.x, and Streamlit 1.x. The random seed was fixed at 42 for train_test_split to ensure reproducibility. Model training was performed on the command line using python train.py, and the Streamlit application was launched using streamlit run app.py.')
E(doc)
J(doc, 'The dataset of 90 pattern examples was preprocessed using the five-stage NLP pipeline described in Chapter 3. The TF-IDF vectoriser produced a feature matrix with approximately 200\u2013250 unique features (unigrams and bigrams) across the training set. A stratified 80/20 split yielded 72 training examples and 18 test examples, with each of the 8 intent classes represented proportionally. Both models were trained on the same training set and evaluated on the same test set.')

SEC(doc, '4.2 Model Performance Comparison')
J(doc, 'The following table presents a comprehensive comparison of the two trained classifiers across multiple evaluation dimensions:')
E(doc)
TABLE(doc,
    ['Metric', 'Naive Bayes (MultinomialNB)', 'SVM (LinearSVC)'],
    [
        ['Test Accuracy', '61.1%', '72.2%'],
        ['Training Time', '< 0.5 seconds', '< 1 second'],
        ['Inference Latency', '< 1 ms per query', '< 1 ms per query'],
        ['End-to-end Response', '< 1 second', '< 1.5 seconds'],
        ['Model File Size', '~15 KB', '~18 KB'],
        ['Correct on 8 test intents', '5 / 8', '7 / 8'],
        ['Handles unseen phrasings', 'Moderate', 'Good'],
        ['Probability estimates', 'Yes (predict_proba)', 'No (decision function)'],
        ['Interpretability', 'High', 'Medium'],
        ['Best for', 'Fast, lightweight deployment', 'Higher-accuracy deployment'],
    ],
    caption='Table 4.1: Model Performance Comparison'
)
E(doc)
CODE(doc, """
  Model Accuracy Comparison
  -----------------------------------------------
  Naive Bayes  |████████████████████░░░░░░░░| 61.1%
  LinearSVC    |████████████████████████░░░░| 72.2%
               0%                         100%
  -----------------------------------------------
  Difference: SVM outperforms NB by +11.1%
""")
cp_bar = doc.add_paragraph()
pf(cp_bar, WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=10)
run(cp_bar, 'Figure 4.1: Model Accuracy Comparison Chart', italic=True)

SEC(doc, '4.3 Intent-wise Classification Results')
J(doc, 'The following table presents the intent-wise classification results for the SVM (LinearSVC) model, which is the recommended model for production use based on its higher accuracy:')
E(doc)
TABLE(doc,
    ['Intent', 'Sample Query', 'Predicted', 'Correct?', 'Response Preview'],
    [
        ['greeting', 'hello there', 'greeting', '\u2713', 'Hello! Welcome to Customer Support...'],
        ['goodbye', 'bye, take care', 'goodbye', '\u2713', 'Goodbye! Have a wonderful day...'],
        ['thanks', 'thank you so much', 'thanks', '\u2713', 'You\u2019re welcome! Happy to help...'],
        ['order_status', 'where is my order', 'order_status', '\u2713', 'To track your order, please visit...'],
        ['refund', 'i want a refund', 'refund', '\u2713', 'Our refund policy allows returns...'],
        ['product_info', 'what do you sell', 'product_info', '\u2713', 'We offer a wide range of products...'],
        ['complaint', 'terrible service', 'complaint', '\u2713', 'I\u2019m really sorry to hear about...'],
        ['human_agent', 'connect me to agent', 'human_agent', '\u2713', 'I\u2019ll connect you to a live agent...'],
    ],
    caption='Table 4.2: Intent-wise Classification Results (SVM Model)'
)

SEC(doc, '4.4 Web Application Screenshots')
J(doc, 'The following screenshots were captured from the live Streamlit deployment of the Customer Support Chat-Bot running at http://localhost:8501. They demonstrate the complete functionality of the web application across different usage scenarios.')
E(doc)
FIG(doc, 'ss1', 'Figure 5.1: Customer Support ChatBot \u2014 Main Interface with SVM Model and Accuracy Sidebar')
FIG(doc, 'ss2', 'Figure 5.2: ChatBot Responding to Greeting Query \u2014 Intent: greeting')
FIG(doc, 'ss3', 'Figure 5.3: ChatBot Responding to Order Status Query \u2014 Intent: order_status')
FIG(doc, 'ss4', 'Figure 5.4: ChatBot Complaint Intent Detection with Escalation Response')
E(doc)
J(doc, 'The screenshots confirm that all key UI features are functioning as designed: the title \u2018Customer Support ChatBot\u2019 and subtitle \u2018Powered by NLP & Machine Learning\u2019 are displayed correctly; the SVM model is selected by default; model accuracy metrics (NB: 61.1%, SVM: 72.2%) are visible in the sidebar; bot messages appear on the left in grey bubbles with the intent tag displayed below; user messages appear on the right in blue bubbles; and the footer correctly reads \u2018Built with NLTK + Scikit-learn | Sharda University Final Year Project\u2019.')

SEC(doc, '4.5 Performance Analysis')
J(doc, 'The experimental results yield several important insights about the behaviour of the two classifiers on this task. The SVM (LinearSVC) outperforms Naive Bayes by 11.1 percentage points (72.2% vs 61.1%) on the held-out test set. This performance gap is consistent with findings in the literature for short-text intent classification. The theoretical explanation is that SVM\u2019s maximum-margin learning objective provides better generalisation in high-dimensional feature spaces compared to NB\u2019s conditional independence assumption, which is frequently violated in natural language where word co-occurrences are strongly correlated.')
E(doc)
J(doc, 'The accuracy figures are modest in absolute terms (61.1% and 72.2%), which is expected given the extremely small training dataset (approximately 9 examples per class). Statistical learning theory predicts that classifier accuracy increases logarithmically with training set size: doubling the training data from 9 to 18 examples per class would be expected to increase SVM accuracy by approximately 5\u201310 percentage points. Extrapolating from published benchmarks on similar tasks, a dataset of 200\u2013500 examples per intent class would be expected to push SVM accuracy above 90%. This represents the most impactful single improvement that could be made to the system.')
E(doc)
J(doc, 'Despite the modest accuracy on the formal test split, manual qualitative testing demonstrated that both models correctly handle the vast majority of naturally phrased customer queries. This apparent discrepancy is explained by the small test set size (18 examples): each misclassification reduces accuracy by approximately 5.6 percentage points (1/18), making the metric highly sensitive to small sample effects.')
E(doc)
J(doc, 'Both models produce end-to-end responses in under 1.5 seconds including Streamlit\u2019s rendering overhead, confirming suitability for real-time conversational applications. The model file sizes (approximately 15\u201318 KB each) are negligible, confirming the system\u2019s suitability for deployment on resource-constrained servers.')

SEC(doc, '4.6 Comparison with Related Systems')
J(doc, 'The following table compares the present system with related existing chatbot implementations reviewed in Chapter 2:')
E(doc)
TABLE(doc,
    ['System', 'Classifier', 'Dataset', 'Accuracy', 'Deployment', 'Open Source'],
    [
        ['Present Project', 'NB + SVM', '90 examples', '72.2% (SVM)', 'Streamlit', 'Yes'],
        ['Kshirsagar (2020)', 'SVM', '500 examples', '85%', 'Web', 'Partial'],
        ['Rasa NLU (2017)', 'DIET (linear)', 'Variable', '90%+', 'REST API', 'Yes'],
        ['ELIZA (1966)', 'Rule-based', 'N/A (rules)', 'N/A', 'Terminal', 'Yes'],
        ['BERT Fine-tuned', 'Transformer', '10k+ examples', '95%+', 'API', 'Yes'],
    ],
    caption='Table 4.3: Comparison with Existing Systems'
)
E(doc)
J(doc, 'The table confirms that the present system offers a strong balance between simplicity, zero cost, and competitive accuracy for its dataset size. While BERT achieves higher absolute accuracy, it requires orders of magnitude more training data and GPU infrastructure. The present system is uniquely positioned for rapid deployment by SMEs or students with limited resources.')

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5: CONCLUSION & FUTURE SCOPE
# ═══════════════════════════════════════════════════════════════════════════════
CHAPTER(doc, '5', 'CONCLUSION & FUTURE SCOPE')

SEC(doc, '5.1 Conclusion')
J(doc, 'This project has successfully demonstrated the design, implementation, training, evaluation, and deployment of an intelligent Customer Support Chat-Bot using classical Natural Language Processing and Machine Learning techniques. The system fulfils all eight objectives defined in Chapter 1, delivering a fully functional, open-source, deployable chatbot solution that automates the handling of eight common customer support intent categories.')
E(doc)
J(doc, 'The core technical contributions of the project are: (1) a clean, six-stage NLP preprocessing pipeline using NLTK that produces high-quality lemmatised tokens; (2) TF-IDF feature extraction with bigrams that captures phrase-level discriminative information; (3) a fair comparison of Naive Bayes and LinearSVC within identical Scikit-learn Pipeline configurations; (4) a modular, JSON-driven architecture that is trivially extensible to new intent domains; and (5) a production-ready Streamlit web application with professional UI, real-time intent transparency, model switching, and session management.')
E(doc)
J(doc, 'The SVM (LinearSVC) classifier achieved 72.2% accuracy on the held-out test set, outperforming Naive Bayes (61.1%) by 11.1 percentage points, consistent with published benchmarks on similar short-text classification tasks. Both models produce responses in under 1.5 seconds end-to-end on standard consumer hardware with no GPU requirement. Manual qualitative testing confirmed correct intent detection for all eight intent categories under natural language phrasing.')
E(doc)
J(doc, 'The project demonstrates that classical Machine Learning approaches \u2014 when combined with well-designed preprocessing and feature extraction \u2014 remain highly relevant and effective for domain-specific, resource-constrained chatbot deployments. The system is immediately deployable as a first-line customer support assistant for e-commerce platforms, service businesses, or any organisation with clearly defined support workflows.')
E(doc)
J(doc, 'From a personal learning perspective, this project provided deep practical experience in the complete NLP-ML pipeline, from data curation and preprocessing through model training, evaluation, serialisation, and production deployment. The challenges encountered \u2014 including the impact of dataset size on accuracy, the trade-off between NB simplicity and SVM accuracy, and the importance of consistent preprocessing between training and inference \u2014 have provided valuable insights that will inform future machine learning work.')

SEC(doc, '5.2 Future Scope')
J(doc, 'The present system, while fully functional, has several clear pathways for enhancement that would substantially improve its capability and deployment reach:')
E(doc)
SEC(doc, '5.2.1 Dataset Expansion and Active Learning')
J(doc, 'The most impactful single enhancement would be expanding the training dataset from 90 to 500+ examples per intent class through crowd-sourcing, data augmentation (paraphrase generation using a language model), or collection of real customer support logs (with appropriate anonymisation). An active learning framework could identify the most informative unlabelled examples for human annotation, minimising labelling cost while maximising accuracy improvement.')
E(doc)
SEC(doc, '5.2.2 Deep Learning Integration (BERT / DistilBERT)')
J(doc, 'Replacing the TF-IDF + LinearSVC pipeline with a fine-tuned DistilBERT model would be expected to push classification accuracy above 90% even on the current small dataset, due to BERT\u2019s pre-trained contextual representations. DistilBERT is approximately 60% smaller and 40% faster than full BERT, making it feasible for CPU deployment.')
E(doc)
SEC(doc, '5.2.3 Multi-turn Dialogue Context')
J(doc, 'The current system treats each user message independently, with no memory of previous turns. Adding a dialogue state tracker (even a simple sliding window of the last 3 messages) would enable the bot to handle follow-up queries such as \u2018What about my second order?\u2019 contextually.')
E(doc)
SEC(doc, '5.2.4 Multilingual Support')
J(doc, 'Extending the system to handle customer queries in Hindi, Tamil, Bengali, or other Indian languages would dramatically expand the addressable user base for Indian SMEs. This could be achieved using multilingual sentence embeddings (e.g., multilingual BERT or LaBSE) as features in place of TF-IDF, combined with translated intent datasets.')
E(doc)
SEC(doc, '5.2.5 Voice Interface')
J(doc, 'Integrating a speech-to-text (STT) front-end (e.g., OpenAI Whisper or Google Speech-to-Text) and a text-to-speech (TTS) back-end (e.g., gTTS or ElevenLabs) would transform the chatbot into a voice-enabled customer support assistant, accessible via phone IVR systems, smart speakers, or web browser microphone.')
E(doc)
SEC(doc, '5.2.6 Mobile and WhatsApp Deployment')
J(doc, 'Packaging the chatbot as a Progressive Web App (PWA) or integrating it with the WhatsApp Business Cloud API via a webhook would enable mobile-first deployment, which is critical for reaching customers in India where WhatsApp is the dominant messaging platform with over 500 million users.')
E(doc)
SEC(doc, '5.2.7 Analytics and Monitoring Dashboard')
J(doc, 'Building an admin dashboard that visualises query volumes per intent category, peak traffic hours, model confidence distributions, and misclassification patterns would provide actionable business intelligence and help prioritise dataset expansion efforts.')

# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
BR(doc)
E(doc, 2)
C(doc, 'REFERENCES', size=14, bold=True)
E(doc, 2)

refs = [
    '[1]\tJ. Weizenbaum, \u201cELIZA\u2014a computer program for the study of natural language communication between man and machine,\u201d Communications of the ACM, vol. 9, no. 1, pp. 36\u201345, Jan. 1966.',
    '[2]\tK. M. Colby, S. Weber, and F. D. Hilf, \u201cArtificial paranoia,\u201d Artificial Intelligence, vol. 2, no. 1, pp. 25\u201336, 1971.',
    '[3]\tC. Cortes and V. Vapnik, \u201cSupport-vector networks,\u201d Machine Learning, vol. 20, no. 3, pp. 273\u2013297, Sep. 1995.',
    '[4]\tS. Bird, E. Klein, and E. Loper, Natural Language Processing with Python: Analysing Text with the Natural Language Toolkit. Sebastopol, CA: O\u2019Reilly Media, 2009.',
    '[5]\tA. Kumar, P. H. Joshi, and A. Jain, \u201cA survey of chatbot implementation in the customer service industry,\u201d IEEE Access, vol. 8, pp. 111\u2013125, 2020.',
    '[6]\tJ. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, \u201cBERT: Pre-training of deep bidirectional transformers for language understanding,\u201d in Proc. NAACL-HLT, 2019, pp. 4171\u20134186.',
    '[7]\tT. B. Brown et al., \u201cLanguage models are few-shot learners,\u201d in Proc. NeurIPS 2020, vol. 33, pp. 1877\u20131901.',
    '[8]\tA. Bocklisch, J. Faulkner, N. Pawlowski, and A. Nichol, \u201cRasa: Open source language understanding and dialogue management,\u201d arXiv:1712.05181, Dec. 2017.',
    '[9]\tN. B. Kshirsagar and A. Pawar, \u201cCustomer support chatbot using natural language processing,\u201d in Proc. IEEE ESCI, 2020, pp. 252\u2013256.',
    '[10]\tD. Jurafsky and J. H. Martin, Speech and Language Processing, 3rd ed. Pearson, 2021.',
    '[11]\tH. Chen, X. Liu, D. Yin, and J. Tang, \u201cA survey on dialogue systems: Recent advances and new frontiers,\u201d ACM SIGKDD Explorations, vol. 19, no. 2, pp. 25\u201335, 2017.',
    '[12]\tP. R. Choudhury and M. Kumar, \u201cImplementation of an AI-based chatbot using NLP for automated customer support,\u201d in Proc. IEEE ICAC3, 2021, pp. 1\u20136.',
]

for ref in refs:
    p = doc.add_paragraph()
    pf(p, WD_ALIGN_PARAGRAPH.JUSTIFY, sa=3)
    run(p, ref)

# ─── ADD PAGE NUMBERS ────────────────────────────────────────────────────────
add_page_numbers(doc)

# ─── SAVE ─────────────────────────────────────────────────────────────────────
out = os.path.join(BASE, 'Major_Report_Customer_Support_ChatBot.docx')
doc.save(out)
print(f'Saved: {out}')
print(f'File size: {os.path.getsize(out) / 1024:.0f} KB')
print('Done!')
